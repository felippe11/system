from io import BytesIO
from typing import List, Dict, Any
import tempfile
import os
import logging

from docx import Document
from docx2pdf import convert as docx2pdf_convert
from transformers import pipeline


# Initialize a lightweight text generation pipeline
# Using t5-small for free summarization/text generation
_model = pipeline("text2text-generation", model="t5-small")


def gerar_texto_relatorio(evento, dados_selecionados: List[str]) -> str:
    """Gera texto de relatório para um evento usando modelo T5.

    Parameters
    ----------
    evento: models.Evento
        Instância do evento.
    dados_selecionados: List[str]
        Lista de campos ou informações selecionadas para compor o relatório.

    Returns
    -------
    str
        Texto gerado pelo modelo com base nos dados fornecidos.
    """
    entrada = f"Evento: {getattr(evento, 'nome', '')}. " + " ".join(dados_selecionados)
    resultado = _model(entrada, max_length=200, do_sample=False)
    return resultado[0]["generated_text"].strip()


def criar_documento_word(texto: str, cabecalho: str = "", rodape: str = "", dados: Dict[str, Any] | None = None) -> BytesIO:
    """Cria um documento Word com cabeçalho, rodapé e corpo de texto.

    Parameters
    ----------
    texto: str
        Corpo principal do relatório.
    cabecalho: str
        Conteúdo do cabeçalho.
    rodape: str
        Conteúdo do rodapé.
    dados: Dict[str, Any] | None
        Dados adicionais para inclusão no documento.

    Returns
    -------
    BytesIO
        Documento Word em memória.
    """
    doc = Document()
    section = doc.sections[0]
    if cabecalho:
        section.header.paragraphs[0].text = cabecalho
    if rodape:
        section.footer.paragraphs[0].text = rodape

    for linha in texto.split("\n"):
        doc.add_paragraph(linha)

    if dados:
        for chave, valor in dados.items():
            doc.add_paragraph(f"{chave}: {valor}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def converter_para_pdf(docx_bytes: BytesIO) -> BytesIO:
    """Converte bytes de um documento Word para PDF.

    Parameters
    ----------
    docx_bytes: BytesIO
        Documento Word em memória.

    Returns
    -------
    BytesIO
        PDF gerado a partir do documento Word.
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
        tmp_docx.write(docx_bytes.getvalue())
        tmp_docx.flush()
        tmp_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        tmp_pdf.close()
        try:
            docx2pdf_convert(tmp_docx.name, tmp_pdf.name)
        except Exception as e:
            logging.error("Falha na conversão com docx2pdf: %s", e)
            try:
                import pypandoc
                pypandoc.convert_file(tmp_docx.name, 'pdf', outputfile=tmp_pdf.name)
            except Exception as e2:
                logging.error("Fallback pypandoc também falhou: %s", e2)
                os.unlink(tmp_docx.name)
                os.unlink(tmp_pdf.name)
                docx_bytes.seek(0)
                return docx_bytes
        with open(tmp_pdf.name, "rb") as f:
            pdf_data = f.read()
    os.unlink(tmp_docx.name)
    os.unlink(tmp_pdf.name)
    return BytesIO(pdf_data)
